"""Build a lean, table/figure-driven Technical Evaluation .docx.

Mostly tables + charts with short connecting text. Embeds the full results
table. Run:  .venv/bin/python docs/technical_evaluation/build_doc.py
"""
from __future__ import annotations

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

CH = "docs/technical_evaluation/charts"
BLUE = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x66, 0x66, 0x66)
RED = RGBColor(0xB0, 0x00, 0x00)

USERS = [10, 50, 100, 200]
AVG = {  # ms
    "coverage":  [1916, 25562, 16358, 42707],
    "datex":     [1624, 11489, 13812, 27541],
    "fused":     [1582, 6343, 12091, 29639],
    "transform": [503, 5993, 12208, 27581],
    "geojson":   [7815, 23939, 37181, 64119],
}
THR = {  # req/s
    "coverage":  ["1.5", "0.56", "1.4", "1.0"],
    "datex":     ["1.6", "1.2", "1.5", "1.3"],
    "fused":     ["1.6", "1.3", "1.5", "1.3"],
    "transform": ["1.6", "1.3", "1.5", "1.3"],
    "geojson":   ["0.13", "0.18", "0.24", "0.29"],
}
SMP = {  # samples
    "coverage":  [60, 300, 600, 1200],
    "datex":     [120, 600, 1200, 2400],
    "fused":     [120, 600, 1200, 2400],
    "transform": [120, 600, 1200, 2400],
    "geojson":   [10, 50, 100, 200],
}
ERR = {  # % at 200 users (0 elsewhere)
    "coverage":  ["0.0", "0.0", "0.0", "0.08"],
    "datex":     ["0.0", "0.0", "0.0", "0.25"],
    "fused":     ["0.0", "0.0", "0.0", "0.12"],
    "transform": ["0.0", "0.0", "0.0", "0.00"],
    "geojson":   ["0.0", "0.0", "0.0", "0.00"],
}
TOTAL_AVG = [1484, 10777, 13783, 31104]
TOTAL_THR = ["5.5", "4.0", "5.2", "4.5"]
TOTAL_ERR = ["0.00", "0.00", "0.00", "0.12"]
TOTAL_SMP = [430, 2150, 4300, 8600]
EP_NAMES = [("coverage", "GET /coverage"), ("datex", "GET /datex"),
            ("fused", "GET /fused"), ("transform", "POST /transform"),
            ("geojson", "GET /geojson")]

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(4)


def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = BLUE
    return h


def P(text="", bold=False, italic=False, size=None, color=None, after=4):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    return p


def B(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    return p


def chart(img, caption, width=6.2):
    doc.add_picture(img, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    P(caption, italic=True, size=8.5, color=GREY, after=8)


def table(headers, rows, style="Light Grid Accent 1", widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = style
    for i, h in enumerate(headers):
        t.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = str(v)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def slot(label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"📸 [ PASTE SCREENSHOT — {label} ]")
    r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = RED
    doc.add_paragraph()


# --------------------------------------------------------------- title
doc.add_heading("Technical Evaluation — Web-based System Benchmark", level=0)
P("DATEX II Road-Weather Adapter · Apache jMeter 5.6.3 · Design Science, Input Session 4",
  italic=True, size=10, color=GREY, after=8)
P("Research question: how do the adapter's response time, throughput and reliability "
  "scale as concurrent consumers increase? (Technical validation addresses system "
  "quality — DeLone & McLean.)", after=8)

# --------------------------------------------------------------- setup
H("1. Experimental setup", 1)
table(["Item", "Value"], [
    ("CPU / cores / freq", "[ from Phase 0 screenshot ]"),
    ("RAM", "[ from Phase 0 screenshot ]"),
    ("Storage", "[ from Phase 0 screenshot ]"),
    ("OS", "[ WSL2 Ubuntu on Windows ]"),
    ("Python", "3.12"),
    ("Web framework", "FastAPI + uvicorn (1 worker, --reload off)"),
    ("Benchmark tool", "Apache jMeter 5.6.3 (Java 21)"),
    ("Network", "localhost loopback"),
    ("System under test", "DATEX II adapter, 1 021 segments, 908 MB demo store"),
])
slot("jMeter GUI with the DATEX-II-Adapter-Benchmark plan")

# --------------------------------------------------------------- method
H("2. Methodology", 1)
B("Three consumer profiles, mixed 60/30/10 % (Normal / Light / Heavy).")
B("Endpoints: GET /coverage · GET /datex · GET /fused · POST /transform · GET /geojson.")
B("1-user warm-up, then 4 load levels: 10 / 50 / 100 / 200 concurrent users.")
B("Same uvicorn config throughout (fair comparison); no response timeout (error = hard failure).")

# --------------------------------------------------------------- results
H("3. Results", 1)

P("Summary (overall, per load level):", bold=True, after=3)
table(["Users", "Samples", "Avg (ms)", "Throughput", "Error %"],
      [(USERS[i], TOTAL_SMP[i], f"{TOTAL_AVG[i]:,}", TOTAL_THR[i] + "/s",
        TOTAL_ERR[i] + "%") for i in range(4)])

P("Full per-endpoint detail:", bold=True, after=3)
rows = []
for u_i in range(4):
    for key, name in EP_NAMES:
        rows.append((USERS[u_i], name, f"{SMP[key][u_i]:,}", f"{AVG[key][u_i]:,}",
                     THR[key][u_i] + "/s", ERR[key][u_i] + "%"))
table(["Users", "Endpoint", "Samples", "Avg (ms)", "Throughput", "Error %"], rows,
      style="Light List Accent 1")

P()
chart(f"{CH}/1_latency_vs_load.png",
      "Figure 1. Average response time vs. concurrent users (log-log). Overall: 1.5 s → 31.1 s.")
chart(f"{CH}/2_throughput_vs_load.png",
      "Figure 2. Throughput vs. concurrent users — flat ceiling ~4.8 req/s.")
chart(f"{CH}/3_percentiles_200users.png",
      "Figure 3. Average / 95th / 99th-percentile latency per endpoint at 200 users.")
chart(f"{CH}/4_error_rate_vs_load.png",
      "Figure 4. Error rate vs. load — 0% to 100 users, 0.12% at 200.")

# --------------------------------------------------------------- findings
H("4. Key findings", 1)
B("Throughput ceiling ≈ 4.8 req/s — flat from 10 to 200 users (20× load, same throughput).")
B("Latency scales with load beyond the ceiling (overall avg 1.5 s → 31.1 s).")
B("Graceful degradation: 0% errors to 100 users; first failures (0.12%) at 200 users.")
B("POST /transform most resilient; GET /geojson (full-network fuse, 4 MB) is the bottleneck.")

# --------------------------------------------------------------- relation
H("5. Relation to previously published results", 1)
P("The README / evaluate.py report an isolated transform latency of ~5.7 ms (p95 ~6.9 ms). "
  "Under concurrent load the same operation's end-to-end latency is 0.5 s (10 users) → 27.6 s "
  "(200 users): the intrinsic compute is fast; the ceiling is a deployment trait (single Python "
  "worker, GIL), not the standardization logic. In the DATEX II / National-Access-Point setting "
  "the per-segment DATEX endpoint stays within second-scale availability at moderate load; the "
  "full-network GeoJSON feed is the operation to manage as consumers grow.")

# --------------------------------------------------------------- bonus
H("6. Performance bug found & fixed via benchmarking", 1)
P("GET /health averaged ~15 s — it re-scanned ~4.3 M rows (908 MB) on every call. Caching the "
  "invariant counts cut steady-state /health from ~13 s to ~0.2 s (~65×).")
chart(f"{CH}/5_health_before_after.png",
      "Figure 5. /health response time before vs. after caching (log scale).", width=4.6)
slot("/health curl output: warm ~8.6 s → cached ~0.2 s")

# --------------------------------------------------------------- threats
H("7. Threats to validity", 1)
B("Single-worker deployment (more workers is an untested lever).")
B("Loopback network — no real latency (upper bound for a co-located consumer).")
B("Shared host — jMeter competes with the server for CPU.")
B("Demo store on WSL2 drvfs is slower than native Linux I/O.")

doc.save("docs/technical_evaluation/Technical_Evaluation.docx")
print("wrote docs/technical_evaluation/Technical_Evaluation.docx")
